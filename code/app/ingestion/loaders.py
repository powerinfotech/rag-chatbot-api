"""문서 로더: .docx (python-docx) / .doc (macOS textutil) → 평문 텍스트."""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path

_BLANK_LINES = re.compile(r"\n[ \t]*\n(?:[ \t]*\n)+")
_TRAILING_WS = re.compile(r"[ \t]+\n")


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _TRAILING_WS.sub("\n", text)
    text = _BLANK_LINES.sub("\n\n", text)
    return text.strip()


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        # python-docx가 본문을 못 뽑는 경우(특이 구조) docx2txt로 폴백.
        import docx2txt

        text = (docx2txt.process(str(path)) or "").strip()
    return text


def _load_doc(path: Path) -> str:
    if shutil.which("textutil") is None:
        raise RuntimeError(
            "legacy .doc 파일 변환에 macOS 'textutil'이 필요하지만 찾을 수 없습니다."
        )
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", "-encoding", "UTF-8", str(path)],
        capture_output=True,
        check=True,
    )
    return result.stdout.decode("utf-8", errors="replace")


def load_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        raw = _load_docx(path)
    elif suffix == ".doc":
        raw = _load_doc(path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {path.suffix}")
    return _normalize(raw)


def iter_document_paths(directory: str | Path) -> list[Path]:
    directory = Path(directory)
    paths = [
        p
        for p in directory.iterdir()
        if p.is_file()
        and p.suffix.lower() in {".doc", ".docx"}
        and not p.name.startswith("~$")
        and p.name != ".DS_Store"
    ]
    return sorted(paths)
